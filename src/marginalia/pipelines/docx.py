"""DOCX pipeline (.docx via python-docx).

Extracts paragraphs + headings + table cells in document order, then
hands off to the shared text indexer. The original docx hierarchy
(Heading 1/2/3 styles) is preserved as a heading prefix `# / ## / ###`
so the indexer can produce heading-anchored sections.

read_segment supports paragraph_start / paragraph_end ranges (1-indexed,
inclusive — only counting non-empty rendered blocks), regex pattern
search, and the generic offset/max_chars chunking over the full body.

Embedded images can be described by the vision profile and inlined near
their source block. Embedded objects and footnotes are skipped.
"""
from __future__ import annotations

import asyncio
import io
import logging
import re
from typing import Any

from marginalia.config import get_settings
from marginalia.pipelines._text_indexer import index_extracted_text
from marginalia.pipelines.base import (
    Pipeline,
    PipelineContext,
    PipelineResult,
    SegmentResult,
)
from marginalia.pipelines.document_vision import (
    DocumentImage,
    answer_document_image_question,
    attach_document_vision_description,
    describe_document_images,
    document_vision_coverage,
    inline_document_image_vision_text,
    persisted_document_image_payload,
    persisted_document_image_segment,
)
from marginalia.pipelines.registry import register_pipeline
from marginalia.storage.base import StorageBackend

log = logging.getLogger(__name__)

# Do not reject DOCX by compressed package size. Media-heavy Word files can
# exceed tens of MB while yielding a small text index; extracted content below
# is still bounded by paragraph and prompt budgets.
MAX_OUTPUT_CHARS = 80_000  # plenty for the LLM prompt
DEFAULT_MAX_CHARS = 8000
_REL_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"


@register_pipeline(
    mimes=(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    exts=(".docx",),
)
class DocxPipeline(Pipeline):
    name = "docx"

    async def run(
        self,
        *,
        ctx: PipelineContext,
        storage: StorageBackend,
    ) -> PipelineResult:
        body_bytes = await self._read_bytes(storage, ctx.storage_key)
        paragraphs = self._parse_paragraphs_from_bytes(body_bytes)
        images = await asyncio.to_thread(_extract_docx_images, body_bytes)
        vision_payload = await describe_document_images(
            settings=get_settings(),
            images=images,
            document_name=ctx.display_name or ctx.storage_key,
            document_kind="docx",
        )
        paragraphs = inline_document_image_vision_text(
            paragraphs,
            vision_payload,
            anchor_key="block",
        )
        full_body = "\n".join(paragraphs)
        indexed_chars = min(len(full_body), MAX_OUTPUT_CHARS)
        body = full_body
        if len(full_body) > MAX_OUTPUT_CHARS:
            body = full_body[:MAX_OUTPUT_CHARS] + "\n[...document truncated for indexing...]"
        coverage = _docx_coverage(
            total_chars=len(full_body),
            indexed_chars=indexed_chars,
            total_paragraphs=len(paragraphs),
        )
        vision_coverage = document_vision_coverage(vision_payload)
        if vision_coverage is not None:
            coverage["document_vision"] = vision_coverage
        result = await index_extracted_text(
            body, ctx, kind="text", coverage=coverage,
        )
        result.description = attach_document_vision_description(
            result.description,
            vision_payload,
        )
        return result

    async def read_segment(
        self,
        *,
        file_row: Any,
        args: dict[str, Any],
        storage: StorageBackend,
    ) -> SegmentResult:
        question = str(args.get("question") or "").strip()
        if question:
            settings = get_settings()
            body = await self._read_bytes(storage, file_row.storage_key)
            images = await asyncio.to_thread(_extract_docx_images, body)
            segment = await answer_document_image_question(
                settings=settings,
                images=images,
                args=args,
                document_name=str(getattr(file_row, "storage_key", "") or ""),
                document_kind="docx",
                mode="docx_image_question",
            )
            if segment.error is None:
                return segment
            fallback = persisted_document_image_segment(
                file_row,
                args,
                mode="docx_image_question",
                warning=segment.error,
            )
            if fallback.error is None:
                return fallback
            return segment

        paragraphs = await self._extract_paragraphs(storage, file_row.storage_key)
        paragraphs = inline_document_image_vision_text(
            paragraphs,
            persisted_document_image_payload(file_row),
            anchor_key="block",
        )
        source_result = self._slice(paragraphs, args, file_row=file_row)
        if source_result.error is None and source_result.text.strip():
            return source_result
        fallback = persisted_document_image_segment(
            file_row,
            args,
            mode="docx_image_vision",
        )
        if fallback.error is None:
            return fallback
        return source_result

    async def read_segment_from_bytes(
        self,
        body: bytes,
        args: dict[str, Any],
        *,
        filename: str | None = None,
    ) -> SegmentResult:
        """Bytes-first variant — used by ArchivePipeline for member peeks."""
        question = str(args.get("question") or "").strip()
        if question:
            settings = get_settings()
            images = await asyncio.to_thread(_extract_docx_images, body)
            return await answer_document_image_question(
                settings=settings,
                images=images,
                args=args,
                document_name=filename or "archive member",
                document_kind="docx",
                mode="docx_image_question",
            )
        try:
            paragraphs = self._parse_paragraphs_from_bytes(body)
        except Exception as exc:  # noqa: BLE001 — python-docx surfaces many
            return SegmentResult(error=f"docx parse failed: {exc}")
        return self._slice(paragraphs, args, file_row=None)

    def _slice(
        self,
        paragraphs: list[str],
        args: dict[str, Any],
        *,
        file_row: Any | None,
    ) -> SegmentResult:
        """Resolve args against this docx body.

        Field priority:
          1. pattern                       → regex search
          2. paragraph_start/_end          → return paragraph range
          3. (default)                     → offset..offset+max_chars chunk
        """
        body = "\n".join(paragraphs)
        total_paragraphs = len(paragraphs)

        offset = max(0, int(args.get("offset") or 0))
        max_chars = int(args.get("max_chars") or DEFAULT_MAX_CHARS)
        if max_chars <= 0:
            max_chars = DEFAULT_MAX_CHARS

        pattern = (args.get("pattern") or "").strip()
        if pattern:
            scope_paragraphs = paragraphs
            paragraph_offset = 0
            ps_raw = args.get("paragraph_start")
            pe_raw = args.get("paragraph_end")
            if ps_raw or pe_raw:
                try:
                    ps = max(1, int(ps_raw)) if ps_raw else 1
                    pe = int(pe_raw) if pe_raw else len(paragraphs)
                except (TypeError, ValueError):
                    return SegmentResult(error="paragraph_start/end must be integers")
                if pe < ps:
                    return SegmentResult(error="paragraph_end must be >= paragraph_start")
                ps = max(1, min(ps, max(1, len(paragraphs))))
                pe = max(ps, min(pe, len(paragraphs)))
                scope_paragraphs = paragraphs[ps - 1: pe]
                paragraph_offset = ps - 1
            return _docx_pattern_search(
                paragraphs=scope_paragraphs, pattern=pattern,
                context_lines=int(args.get("context_lines") or 2),
                max_matches=int(args.get("max_matches") or 20),
                match_offset=max(0, int(args.get("match_offset") or 0)),
                paragraph_offset=paragraph_offset,
                total_paragraphs_full=len(paragraphs),
            )

        if any(args.get(key) for key in ("section_id", "heading", "line_start", "line_end")):
            from marginalia.pipelines.text import TextPipeline

            return TextPipeline()._slice(body=body, args=args, file_row=file_row)

        para_start = args.get("paragraph_start")
        para_end = args.get("paragraph_end")
        if para_start:
            try:
                ps = int(para_start)
            except (TypeError, ValueError):
                return SegmentResult(error="paragraph_start must be an integer")
            try:
                pe = int(para_end) if para_end else ps
            except (TypeError, ValueError):
                return SegmentResult(error="paragraph_end must be an integer")
            if total_paragraphs == 0:
                return SegmentResult(error="docx has no paragraphs")
            ps = max(1, min(ps, total_paragraphs))
            pe = max(ps, min(pe, total_paragraphs))
            slab = "\n".join(paragraphs[ps - 1: pe])
            return _clamp(
                slab, offset, max_chars,
                extras={
                    "paragraph_start": ps,
                    "paragraph_end": pe,
                    "total_paragraphs": total_paragraphs,
                },
            )

        # Compute paragraph range from char offset so footnotes can
        # deep-link even when the LLM reads by offset rather than
        # paragraph_start/paragraph_end.
        para_start = body[:offset].count("\n") + 1
        chunk_for_range = body[offset: offset + max_chars]
        para_end = para_start + chunk_for_range.count("\n")
        return _clamp(
            body, offset, max_chars,
            extras={
                "total_paragraphs": total_paragraphs,
                "paragraph_start": para_start,
                "paragraph_end": para_end,
            },
        )

    @classmethod
    async def _extract_paragraphs(
        cls,
        storage: StorageBackend, key: str,
    ) -> list[str]:
        try:
            from docx import Document  # type: ignore  # noqa: F401 — keeps import-error early
        except ImportError as exc:
            raise RuntimeError(
                "docx pipeline needs python-docx; "
                "`pip install python-docx`"
            ) from exc

        buf = bytearray()
        async for chunk in storage.get(key):
            buf.extend(chunk)
        return cls._parse_paragraphs_from_bytes(bytes(buf))

    @staticmethod
    async def _read_bytes(storage: StorageBackend, key: str) -> bytes:
        buf = bytearray()
        async for chunk in storage.get(key):
            buf.extend(chunk)
        return bytes(buf)

    @staticmethod
    def _parse_paragraphs_from_bytes(body: bytes) -> list[str]:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "docx pipeline needs python-docx; "
                "`pip install python-docx`"
            ) from exc
        doc = Document(io.BytesIO(body))
        out: list[str] = []
        for block in _iter_block_items(doc):
            line = _render_block(block)
            if line:
                out.append(line)
        return out


def _extract_docx_images(body: bytes) -> list[DocumentImage]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "docx pipeline needs python-docx; `pip install python-docx`"
        ) from exc

    try:
        doc = Document(io.BytesIO(body))
    except Exception:
        # Images are supplemental vision fodder: a package python-docx cannot
        # re-open (corrupt zip, odd relationships) must not fail the ingest —
        # the text extraction has already produced the primary content.
        log.warning("docx image extraction skipped: package unreadable",
                    exc_info=True)
        return []
    images: list[DocumentImage] = []
    seen: set[tuple[int, str]] = set()
    rendered_block_no = 0
    for block_no, block in enumerate(_iter_block_items(doc), start=1):
        context = _render_block(block)
        if context:
            rendered_block_no += 1
        element = getattr(block, "_element", None)
        if element is None:
            continue
        anchor_block = max(1, rendered_block_no)
        for rid in _embedded_image_rids(element):
            key = (block_no, rid)
            if key in seen:
                continue
            seen.add(key)
            part = doc.part.related_parts.get(rid)
            blob = getattr(part, "blob", None)
            if not isinstance(blob, bytes) or not blob:
                continue
            image_no = len(images) + 1
            images.append(
                DocumentImage(
                    image_id=f"docx-img-{image_no}",
                    label=f"DOCX image {image_no}",
                    image_bytes=blob,
                    media_type=getattr(part, "content_type", None),
                    anchor={
                        "unit": "blocks",
                        "block": anchor_block,
                        "source_block": block_no,
                        "relationship_id": rid,
                    },
                    context=context,
                    filename=str(getattr(part, "partname", "") or ""),
                )
            )
    return images


def _embedded_image_rids(element: Any) -> list[str]:
    rids: list[str] = []
    for child in element.iter():
        if not str(getattr(child, "tag", "")).endswith("}blip"):
            continue
        rid = child.get(_REL_EMBED)
        if rid:
            rids.append(str(rid))
    return rids


def _iter_block_items(doc: Any):
    """Yield paragraphs and tables in document order.

    python-docx exposes doc.paragraphs and doc.tables as separate lists, so
    walk the underlying body XML to recover order.
    """
    from docx.oxml.ns import qn  # type: ignore
    from docx.table import Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _render_block(block: Any) -> str:
    cls_name = type(block).__name__
    if cls_name == "Paragraph":
        text = (block.text or "").strip()
        if not text:
            return ""
        style = (getattr(block.style, "name", "") or "").strip()
        if style.startswith("Heading 1"):
            return f"# {text}"
        if style.startswith("Heading 2"):
            return f"## {text}"
        if style.startswith("Heading 3"):
            return f"### {text}"
        if style.startswith("Heading"):
            return f"#### {text}"
        return text
    if cls_name == "Table":
        rows: list[str] = []
        for row in block.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    return ""


def _docx_coverage(
    *, total_chars: int, indexed_chars: int, total_paragraphs: int,
) -> dict[str, Any]:
    indexed_partial = indexed_chars < total_chars
    return {
        "unit": "characters",
        "source_mode": "docx_extracted_text",
        "total_units": total_chars,
        "indexed_units": indexed_chars,
        "total_chars": total_chars,
        "indexed_chars": indexed_chars,
        "total_paragraphs": total_paragraphs,
        "indexed_partial": indexed_partial,
        "partial_reasons": ["prompt_text_cap"] if indexed_partial else [],
        "max_index_chars": MAX_OUTPUT_CHARS,
        "chunked": False,
        "chunk_count": 1,
        "text_truncated": indexed_partial,
    }


# ---- read_segment helpers --------------------------------------------------

def _clamp(
    text: str, offset: int, max_chars: int,
    *, extras: dict[str, Any] | None = None,
) -> SegmentResult:
    extras = dict(extras or {})
    total = len(text)
    chunk = text[offset: offset + max_chars]
    truncated = (offset + len(chunk)) < total
    extras.update({
        "offset": offset,
        "char_count": len(chunk),
        "total_chars": total,
        "truncated": truncated,
    })
    if truncated:
        extras["next_offset"] = offset + len(chunk)
    if not chunk:
        return SegmentResult(text="", error="empty result", extras=extras)
    return SegmentResult(text=chunk, extras=extras)


def _docx_pattern_search(
    *, paragraphs: list[str], pattern: str,
    context_lines: int, max_matches: int,
    match_offset: int = 0, paragraph_offset: int = 0,
    total_paragraphs_full: int | None = None,
) -> SegmentResult:
    try:
        rx = re.compile(pattern, re.IGNORECASE | re.MULTILINE)
    except re.error as exc:
        return SegmentResult(error=f"invalid regex: {exc}")

    full_total = (
        total_paragraphs_full if total_paragraphs_full is not None
        else len(paragraphs)
    )

    all_hits: list[dict[str, Any]] = []
    for i, para in enumerate(paragraphs, start=1):
        if not para:
            continue
        for m in rx.finditer(para):
            s = max(0, i - 1 - context_lines)
            e = min(len(paragraphs), i + context_lines)
            all_hits.append({
                "paragraph": i + paragraph_offset,
                "match": m.group(0)[:200],
                "context": "\n".join(paragraphs[s:e]),
            })

    total = len(all_hits)
    hits = all_hits[match_offset: match_offset + max_matches]
    has_more = (match_offset + len(hits)) < total

    extras: dict[str, Any] = {
        "pattern": pattern,
        "match_count": len(hits),
        "total_matches": total,
        "match_offset": match_offset,
        "has_more": has_more,
        "hits": hits,
        "total_paragraphs": full_total,
    }
    if has_more:
        extras["next_match_offset"] = match_offset + len(hits)

    if not hits:
        if match_offset and total:
            err = f"match_offset {match_offset} exceeds total_matches {total}"
        else:
            err = "no matches"
        return SegmentResult(text="", error=err, extras=extras)

    rendered = "\n\n".join(
        f"[¶{h['paragraph']}] {h['match']}\n  ┊ {h['context']}"
        for h in hits
    )
    return SegmentResult(text=rendered, extras=extras)
